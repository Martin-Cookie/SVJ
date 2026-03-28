from __future__ import annotations

"""
Parse a Word (.docx) ballot template to extract voting items.
Looks for patterns like:
  "BOD 1: ..." or "1. ..." or "1) ..."
"""
import logging
import re

from docx import Document

logger = logging.getLogger(__name__)


_CZECH_MONTHS = {
    "ledna": 1, "února": 2, "března": 3, "dubna": 4,
    "května": 5, "června": 6, "července": 7, "srpna": 8,
    "září": 9, "října": 10, "listopadu": 11, "prosince": 12,
}

# Czech ordinal words → number
_CZECH_ORDINALS = {
    "první": 1, "druhý": 2, "druhá": 2, "třetí": 3, "čtvrtý": 4, "čtvrtá": 4,
    "pátý": 5, "pátá": 5, "šestý": 6, "šestá": 6, "sedmý": 7, "sedmá": 7,
    "osmý": 8, "osmá": 8, "devátý": 9, "devátá": 9, "desátý": 10, "desátá": 10,
}

# Date-like text after a number — used to filter out false positives like "19. ledna 2026"
_DATE_AFTER_NUM = re.compile(
    r"^\s*(?:" + "|".join(_CZECH_MONTHS) + r")\s+\d{4}",
    re.IGNORECASE,
)

# Footer/closing patterns — stop capturing description when these appear after voting items
_FOOTER_PATTERN = re.compile(
    r"(?:"
    r"^V\s+\w+\s+dne\b"                    # "V Praze dne ...", "V Brně dne ..."
    r"|^Dne\s+\d"                           # "Dne 19. ledna 2026"
    r"|^Za\s+výbor\b"                       # "Za výbor SVJ"
    r"|^Předseda\s"                         # "Předseda výboru"
    r"|^Podpis[:\s]"                        # "Podpis:", "Podpisy členů"
    r"|^Příloha\b"                          # "Příloha č. 1"
    r"|^Přílohy\b"                          # "Přílohy:"
    r"|^Výsledky\s+hlasování\b"            # "Výsledky hlasování budou"
    r"|^Výsledek\s+hlasování\b"            # "Výsledek hlasování bude"
    r"|^Hlasovací\s+lístky?\s"             # "Hlasovací lístek odevzdejte"
    r"|^Lístek\s+odevzdejte\b"             # "Lístek odevzdejte"
    r"|^Odevzdejte\b"                       # "Odevzdejte do ..."
    r"|^Vyplněný\s+lístek\b"               # "Vyplněný lístek odevzdejte"
    r"|^Poznámka[:\s]"                      # "Poznámka:"
    r"|^\*{3,}"                             # "***" separators
    r"|^-{3,}"                              # "---" separators
    r"|^_{3,}"                              # "___" separators
    r")",
    re.IGNORECASE,
)


def extract_voting_items(docx_path: str) -> list[dict]:
    doc = Document(docx_path)
    items = []
    current_item = None

    # Pattern: Czech ordinal + "bod hlasování" + separator + title
    # e.g. "První bod hlasování – provoz FVE"
    ordinal_words = "|".join(_CZECH_ORDINALS.keys())
    ordinal_pattern = re.compile(
        r"^\s*(" + ordinal_words + r")\s+bod\s+hlasování\s*[–\-:.]\s*(.+)",
        re.IGNORECASE,
    )

    # Numeric patterns
    numeric_patterns = [
        re.compile(r"^\s*BOD\s+(\d+)\s*[:.]\s*(.+)", re.IGNORECASE),
        re.compile(r"^\s*(\d+)\s*[.)]\s*(.+)"),
        re.compile(r"^\s*(\d+)\s*[:.]\s*(.+)"),
    ]

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        matched = False

        # Try Czech ordinal pattern first
        m = ordinal_pattern.match(text)
        if m:
            if current_item:
                items.append(current_item)
            word = m.group(1).lower()
            order = _CZECH_ORDINALS.get(word, len(items) + 1)
            title = m.group(2).strip()
            current_item = {
                "order": order,
                "title": title,
                "description": "",
            }
            matched = True

        # Try numeric patterns
        if not matched:
            for pattern in numeric_patterns:
                match = pattern.match(text)
                if match:
                    title_part = match.group(2).strip()
                    # Skip date false positives (e.g. "19. ledna 2026")
                    if _DATE_AFTER_NUM.match(title_part):
                        break
                    if current_item:
                        items.append(current_item)
                    order = int(match.group(1))
                    current_item = {
                        "order": order,
                        "title": title_part,
                        "description": "",
                    }
                    matched = True
                    break

        if not matched and current_item:
            # Skip table-like content and checkbox markers
            if text.startswith(("SOUHLASÍM", "NESOUHLASÍM", "☐", "□")):
                continue
            # Stop capturing description when footer/closing text is detected
            if _FOOTER_PATTERN.search(text):
                break
            if current_item["description"]:
                current_item["description"] += "\n"
            current_item["description"] += text

    if current_item:
        items.append(current_item)

    return items


def _parse_czech_date(text: str) -> str | None:
    """Try to parse a Czech date from text, return ISO YYYY-MM-DD or None."""
    # DD.MM.YYYY
    m = re.search(r"(\d{1,2})\.\s*(\d{1,2})\.\s*(\d{4})", text)
    if m:
        d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
        return f"{y}-{mo:02d}-{d:02d}"
    # DD. měsíce YYYY  (e.g. "19. ledna 2026")
    m = re.search(
        r"(\d{1,2})\.\s*(" + "|".join(_CZECH_MONTHS) + r")\s+(\d{4})",
        text, re.IGNORECASE,
    )
    if m:
        d = int(m.group(1))
        mo = _CZECH_MONTHS[m.group(2).lower()]
        y = int(m.group(3))
        return f"{y}-{mo:02d}-{d:02d}"
    return None


def extract_voting_metadata(docx_path: str) -> dict:
    """Extract voting metadata (title, description, dates) from Word document."""
    doc = Document(docx_path)
    paragraphs = [(p.text.strip(), p.style.name if p.style else "") for p in doc.paragraphs]

    result = {"title": None, "description": None, "start_date": None, "end_date": None}

    # --- Title ---
    # Build title from "per rollam" / "ROZHODOVÁNÍ" paragraph + following date line
    per_rollam_pattern = re.compile(
        r"(hlasování\s+per\s+rollam|rozhodování\s+per\s+rollam|per\s+rollam)",
        re.IGNORECASE,
    )
    title_idx = None

    # 1. Document properties
    try:
        if doc.core_properties.title and doc.core_properties.title.strip():
            result["title"] = doc.core_properties.title.strip()
    except Exception:
        logger.debug("Nelze přečíst metadata dokumentu %s", docx_path, exc_info=True)

    if not result["title"]:
        # 2. First Heading 1 or Title style
        for i, (text, style) in enumerate(paragraphs):
            if text and style in ("Heading 1", "Title"):
                result["title"] = text
                title_idx = i
                break

    if not result["title"]:
        # 3. Paragraph matching "per rollam" / "rozhodování per rollam"
        for i, (text, _) in enumerate(paragraphs):
            if text and per_rollam_pattern.search(text):
                result["title"] = text
                title_idx = i
                break

    # If title found and next paragraph is a short date-like line, merge it
    if result["title"] and title_idx is not None:
        next_idx = title_idx + 1
        while next_idx < len(paragraphs) and not paragraphs[next_idx][0]:
            next_idx += 1
        if next_idx < len(paragraphs):
            next_text = paragraphs[next_idx][0]
            if next_text and len(next_text) <= 80 and _parse_czech_date(next_text):
                result["title"] = result["title"] + " " + next_text
                title_idx = next_idx  # advance past the merged line

    if not result["title"]:
        # 4. First short non-empty paragraph (max 150 chars)
        for i, (text, _) in enumerate(paragraphs):
            if text and len(text) <= 150:
                result["title"] = text
                title_idx = i
                break

    # --- Dates ---
    full_text = "\n".join(text for text, _ in paragraphs)

    # Start date: "od|zahájení|začátek|vyhlášené" + date
    start_patterns = [
        r"(?:od|zahájení|začátek)[:\s]+(\d{1,2}\.\s*\d{1,2}\.\s*\d{4})",
        r"(?:od|zahájení|začátek)[:\s]+(\d{1,2}\.\s*(?:" + "|".join(_CZECH_MONTHS) + r")\s+\d{4})",
        r"(?:vyhlášen[áéo])\s+(\d{1,2}\.\s*(?:" + "|".join(_CZECH_MONTHS) + r")\s+\d{4})",
        r"(?:vyhlášen[áéo])\s+(\d{1,2}\.\s*\d{1,2}\.\s*\d{4})",
    ]
    for pat in start_patterns:
        m = re.search(pat, full_text, re.IGNORECASE)
        if m:
            result["start_date"] = _parse_czech_date(m.group(1))
            if result["start_date"]:
                break

    # End date: "do|ukončení|konec|odevzdání" + date
    end_patterns = [
        r"(?:do|ukončení|konec|odevzdání)[:\s]+.*?(\d{1,2}\.\s*\d{1,2}\.\s*\d{4})",
        r"(?:do|ukončení|konec|odevzdání)[:\s]+.*?(\d{1,2}\.\s*(?:" + "|".join(_CZECH_MONTHS) + r")\s+\d{4})",
        r"(?:lhůta|termín).*?(?:do|je)\s+(\d{1,2}\.\s*(?:" + "|".join(_CZECH_MONTHS) + r")\s+\d{4})",
        r"(?:lhůta|termín).*?(?:do|je)\s+(\d{1,2}\.\s*\d{1,2}\.\s*\d{4})",
    ]
    for pat in end_patterns:
        m = re.search(pat, full_text, re.IGNORECASE)
        if m:
            result["end_date"] = _parse_czech_date(m.group(1))
            if result["end_date"]:
                break

    # --- Description: text between title and first voting item ---
    if title_idx is not None:
        ordinal_words = "|".join(_CZECH_ORDINALS.keys())
        bod_patterns = [
            re.compile(r"^\s*BOD\s+\d+", re.IGNORECASE),
            re.compile(r"^\s*(?:" + ordinal_words + r")\s+bod\s+hlasování", re.IGNORECASE),
            re.compile(r"^\s*\d+\s*[.)]\s*\S"),
        ]
        bod_idx = None
        for i in range(title_idx + 1, len(paragraphs)):
            text, _ = paragraphs[i]
            if not text:
                continue
            if any(p.match(text) for p in bod_patterns):
                # Exclude date false positives
                if not _parse_czech_date(text):
                    bod_idx = i
                    break

        if bod_idx is not None:
            skip_re = re.compile(
                r"^(SOUHLASÍM|NESOUHLASÍM|☐|□|se\s+sídlem|zapsané\s+ve|IČO)",
                re.IGNORECASE,
            )
            desc_parts = []
            for i in range(title_idx + 1, bod_idx):
                text, _ = paragraphs[i]
                if not text or skip_re.match(text):
                    continue
                # Skip the date line that was already merged into title
                if _parse_czech_date(text) and len(text) <= 80:
                    continue
                desc_parts.append(text)
            if desc_parts:
                result["description"] = "\n".join(desc_parts)

    return result


def extract_full_text(docx_path: str) -> str:
    doc = Document(docx_path)
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text.strip())
    return "\n".join(texts)
